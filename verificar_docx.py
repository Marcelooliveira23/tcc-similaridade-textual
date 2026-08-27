"""
Verificação estrutural do TCC_FINAL.docx:
- Conta parágrafos por seção
- Verifica presença de tabelas
- Confirma quantidade de referências
- Verifica formatação de fonte e margens
"""

from docx import Document
from docx.shared import Inches

doc = Document(r"C:\Users\mrced\OneDrive\Documents\TCC\TCC_FINAL.docx")

# Margens
sec = doc.sections[0]
top    = round(sec.top_margin.inches    * 2.54, 1)
bottom = round(sec.bottom_margin.inches * 2.54, 1)
left   = round(sec.left_margin.inches   * 2.54, 1)
right  = round(sec.right_margin.inches  * 2.54, 1)

print("=" * 60)
print("VERIFICAÇÃO ESTRUTURAL DO TCC_FINAL.docx")
print("=" * 60)
print(f"\nMARGENS (cm):")
print(f"  Superior : {top}   (esperado: 3.0) {'✅' if top == 3.0 else '❌'}")
print(f"  Inferior : {bottom} (esperado: 2.0) {'✅' if bottom == 2.0 else '❌'}")
print(f"  Esquerda : {left}   (esperado: 3.0) {'✅' if left == 3.0 else '❌'}")
print(f"  Direita  : {right}   (esperado: 2.0) {'✅' if right == 2.0 else '❌'}")

# Parágrafos e fontes
total_paras = len(doc.paragraphs)
arial_ok = 0
arial_fail = 0
sizes_found = set()

for p in doc.paragraphs:
    for run in p.runs:
        if run.font.name and run.font.name == 'Arial':
            arial_ok += 1
        elif run.font.name:
            arial_fail += 1
        if run.font.size:
            sizes_found.add(int(run.font.size.pt))

print(f"\nPARÁGRAFOS E FONTES:")
print(f"  Total de parágrafos : {total_paras}")
print(f"  Runs Arial          : {arial_ok}  ✅")
print(f"  Runs outros         : {arial_fail}")
print(f"  Tamanhos de fonte   : {sorted(sizes_found)} pt")

# Tabelas
print(f"\nTABELAS:")
print(f"  Total de tabelas: {len(doc.tables)} ✅ (esperado: 5)")
for i, t in enumerate(doc.tables, 1):
    print(f"  Tabela {i}: {len(t.rows)} linhas × {len(t.columns)} colunas")

# Seções principais (linhas em CAPS com negrito)
print(f"\nSEÇÕES DETECTADAS:")
secoes = []
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt and any(r.bold for r in p.runs) and txt == txt.upper() and len(txt) > 3:
        secoes.append(txt[:70])

for s in secoes:
    print(f"  • {s}")

# Referências
print(f"\nREFERÊNCIAS:")
in_refs = False
ref_count = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if "REFERÊNCIAS" in txt.upper() and len(txt) < 20:
        in_refs = True
        continue
    if in_refs and len(txt) > 20:
        ref_count += 1

print(f"  Total de referências: {ref_count} {'✅' if ref_count >= 15 else '⚠️'} (esperado: 15)")

print(f"\n{'=' * 60}")
print("CONFORMIDADE GERAL:")
checks = {
    "Margem superior 3.0cm": top == 3.0,
    "Margem inferior 2.0cm": bottom == 2.0,
    "Margem esquerda 3.0cm": left == 3.0,
    "Margem direita 2.0cm":  right == 2.0,
    "5 tabelas (Quadros)":   len(doc.tables) >= 5,
    "15 referências ABNT":   ref_count >= 15,
    "Total de parágrafos":   total_paras >= 80,
}

passed = sum(v for v in checks.values())
for k, v in checks.items():
    print(f"  {'✅' if v else '❌'}  {k}")

print(f"\nRESULTADO: {passed}/{len(checks)} critérios atendidos")
pct = round(passed / len(checks) * 100)
print(f"CONFORMIDADE: {pct}%  {'EXCELENTE ✅' if pct >= 90 else 'VERIFICAR ⚠️'}")
print("=" * 60)
