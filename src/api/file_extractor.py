"""
Extrator universal de texto a partir de arquivos enviados via upload.
Suporta: .txt, .pdf, .docx, .doc, .odt, .rtf, .csv, .md, .html/.htm
"""

import io
import os


# Extensões suportadas
SUPPORTED_EXTENSIONS = {
    ".txt", ".text",
    ".pdf",
    ".docx",
    ".doc",
    ".odt",
    ".rtf",
    ".xlsx", ".xlsm",
    ".csv",
    ".md", ".markdown",
    ".json", ".xml", ".yml", ".yaml", ".sql",
    ".py", ".java", ".c", ".cpp", ".h", ".hpp", ".cs", ".js", ".ts", ".html", ".htm", ".css", ".m",
    ".html", ".htm",
}


def extract_text(file_storage) -> tuple[str, str | None]:
    """
    Extrai texto de um FileStorage (Flask) com base na extensão.

    Args:
        file_storage: objeto FileStorage do Flask (request.files)

    Returns:
        Tupla (texto_extraido, mensagem_de_erro_ou_None)
    """
    filename = file_storage.filename or ""
    raw = file_storage.read()
    return extract_text_from_bytes(raw, filename)


def extract_text_from_bytes(raw: bytes, filename: str) -> tuple[str, str | None]:
    """Extrai texto a partir de bytes e nome de arquivo.

    Essa função permite paralelizar extração de múltiplos arquivos sem depender
    de objetos de upload/stream compartilhados.
    """
    ext = os.path.splitext(filename or "")[1].lower()

    if not ext:
        ext = ".txt"

    if ext in (
        ".txt",
        ".text",
        ".csv",
        ".md",
        ".markdown",
        ".json",
        ".xml",
        ".yml",
        ".yaml",
        ".sql",
        ".py",
        ".java",
        ".c",
        ".cpp",
        ".h",
        ".hpp",
        ".cs",
        ".js",
        ".ts",
        ".css",
        ".m",
    ):
        return _extract_plain_text(raw)

    if ext in (".xlsx", ".xlsm"):
        return _extract_excel(raw)

    if ext == ".pdf":
        return _extract_pdf(raw)

    if ext == ".docx":
        return _extract_docx(raw)

    if ext == ".doc":
        return _extract_doc(raw)

    if ext == ".odt":
        return _extract_odt(raw)

    if ext == ".rtf":
        return _extract_rtf(raw)

    if ext in (".html", ".htm"):
        return _extract_html(raw)

    # Tipo desconhecido — tenta como texto
    text, err = _extract_plain_text(raw)
    if err or not text.strip():
        return "", (
            f"Formato '{ext}' não suportado diretamente. "
            f"Formatos aceitos: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return text, None


# ─── Extratores específicos ───────────────────────────────────────────────────

def _extract_plain_text(raw: bytes) -> tuple[str, str | None]:
    """TXT, CSV, MD — leitura direta."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return raw.decode(enc), None
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("latin-1", errors="replace"), None


def _extract_pdf(raw: bytes) -> tuple[str, str | None]:
    """PDF — usa pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "Biblioteca pypdf não instalada. Execute: pip install pypdf"

    try:
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        text = "\n".join(pages)
        if not text.strip():
            return "", "PDF parece ser escaneado (imagem). Não é possível extrair texto automaticamente."
        return text, None
    except Exception as exc:
        return "", f"Erro ao ler PDF: {exc}"


def _extract_docx(raw: bytes) -> tuple[str, str | None]:
    """DOCX — usa python-docx."""
    try:
        from docx import Document
    except ImportError:
        return "", "Biblioteca python-docx não instalada. Execute: pip install python-docx"

    try:
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Incluir texto de tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs), None
    except Exception as exc:
        return "", f"Erro ao ler DOCX: {exc}"


def _extract_doc(raw: bytes) -> tuple[str, str | None]:
    """DOC (formato legado Word) — tenta extração básica."""
    # Tentar como UTF-16 (formato interno de muitos DOC)
    for enc in ("utf-16", "utf-8", "latin-1"):
        try:
            text = raw.decode(enc, errors="ignore")
            # Filtrar caracteres de controle, manter apenas texto legível
            clean = "".join(c for c in text if c.isprintable() or c in "\n\r\t ")
            clean = clean.strip()
            if len(clean) > 50:
                return clean, None
        except Exception:
            continue

    return "", (
        "Arquivo .doc (Word 97-2003) não pôde ser lido automaticamente. "
        "Salve como .docx ou .txt no Word e tente novamente."
    )


def _extract_odt(raw: bytes) -> tuple[str, str | None]:
    """ODT (LibreOffice) — usa odfpy."""
    try:
        from odf.opendocument import load as odf_load
        from odf.text import P
        from odf import teletype
    except ImportError:
        return "", "Biblioteca odfpy não instalada. Execute: pip install odfpy"

    try:
        doc = odf_load(io.BytesIO(raw))
        paragraphs = []
        for elem in doc.text.getElementsByType(P):
            t = teletype.extractText(elem)
            if t.strip():
                paragraphs.append(t)
        return "\n".join(paragraphs), None
    except Exception as exc:
        return "", f"Erro ao ler ODT: {exc}"


def _extract_rtf(raw: bytes) -> tuple[str, str | None]:
    """RTF — usa striprtf."""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        return "", "Biblioteca striprtf não instalada. Execute: pip install striprtf"

    try:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                rtf_str = raw.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            rtf_str = raw.decode("latin-1", errors="replace")

        text = rtf_to_text(rtf_str)
        return text, None
    except Exception as exc:
        return "", f"Erro ao ler RTF: {exc}"


def _extract_html(raw: bytes) -> tuple[str, str | None]:
    """HTML — extrai texto removendo tags."""
    for enc in ("utf-8", "latin-1"):
        try:
            html = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        html = raw.decode("latin-1", errors="replace")

    # Remover tags HTML simples via regex
    import re
    # Remover scripts e estilos
    html = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html, flags=re.DOTALL | re.IGNORECASE)
    # Remover tags
    text = re.sub(r"<[^>]+>", " ", html)
    # Decodificar entidades comuns
    entities = {
        "&amp;": "&", "&lt;": "<", "&gt;": ">",
        "&nbsp;": " ", "&quot;": '"', "&#39;": "'",
    }
    for ent, char in entities.items():
        text = text.replace(ent, char)
    # Compactar espaços
    text = re.sub(r"\s+", " ", text).strip()
    return text, None


def _extract_excel(raw: bytes) -> tuple[str, str | None]:
    """XLSX/XLSM — extrai conteúdo textual de células."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "", "Biblioteca openpyxl não instalada. Execute: pip install openpyxl"

    try:
        workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        lines: list[str] = []

        for sheet in workbook.worksheets:
            lines.append(f"[sheet] {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if values:
                    lines.append(" | ".join(values))

        text = "\n".join(lines).strip()
        if not text:
            return "", "Planilha sem conteúdo textual para extração."
        return text, None
    except Exception as exc:
        return "", f"Erro ao ler arquivo Excel: {exc}"
