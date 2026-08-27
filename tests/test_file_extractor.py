"""Testes de extração de texto multi-formato."""

import io
import pytest

# ─── helpers para criar FileStorage simulado ─────────────────────────────────

class FakeFileStorage:
    """Simula flask.FileStorage."""
    def __init__(self, content: bytes, filename: str):
        self._content = content
        self.filename = filename
        self._stream = io.BytesIO(content)

    def read(self):
        return self._content


def make_txt(text: str, filename="doc.txt") -> FakeFileStorage:
    return FakeFileStorage(text.encode("utf-8"), filename)


# ─── TXT ─────────────────────────────────────────────────────────────────────

def test_extract_txt():
    from src.api.file_extractor import extract_text
    fs = make_txt("O rato roeu a roupa do rei.", "teste.txt")
    text, err = extract_text(fs)
    assert err is None
    assert "rato" in text


def test_extract_md():
    from src.api.file_extractor import extract_text
    content = "# Título\n\nConteúdo do documento markdown."
    fs = make_txt(content, "readme.md")
    text, err = extract_text(fs)
    assert err is None
    assert "markdown" in text


def test_extract_csv():
    from src.api.file_extractor import extract_text
    content = "nome,valor\nAlpha,10\nBeta,20"
    fs = make_txt(content, "data.csv")
    text, err = extract_text(fs)
    assert err is None
    assert "Alpha" in text


def test_extract_html():
    from src.api.file_extractor import extract_text
    html = "<html><body><h1>Título</h1><p>Parágrafo de texto.</p></body></html>"
    fs = make_txt(html, "page.html")
    text, err = extract_text(fs)
    assert err is None
    assert "Parágrafo" in text
    assert "<p>" not in text  # tags removidas


# ─── PDF ─────────────────────────────────────────────────────────────────────

def test_extract_pdf():
    from src.api.file_extractor import extract_text
    from pypdf import PdfWriter
    writer = PdfWriter()
    page = writer.add_blank_page(200, 200)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    # PDF em branco — espera mensagem de aviso, não exceção
    fs = FakeFileStorage(buf.read(), "test.pdf")
    text, err = extract_text(fs)
    # PDF vazio retorna erro de "escaneado" ou texto vazio, mas não exceção
    assert isinstance(text, str)


# ─── DOCX ────────────────────────────────────────────────────────────────────

def test_extract_docx():
    from src.api.file_extractor import extract_text
    from docx import Document
    doc = Document()
    doc.add_paragraph("Texto de teste no documento Word.")
    doc.add_paragraph("Segunda linha de conteúdo.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fs = FakeFileStorage(buf.read(), "documento.docx")
    text, err = extract_text(fs)
    assert err is None
    assert "Texto de teste" in text
    assert "Segunda linha" in text


def test_extract_docx_com_tabela():
    from src.api.file_extractor import extract_text
    from docx import Document
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Célula A1"
    table.rows[0].cells[1].text = "Célula B1"
    table.rows[1].cells[0].text = "Célula A2"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fs = FakeFileStorage(buf.read(), "tabela.docx")
    text, err = extract_text(fs)
    assert err is None
    assert "Célula A1" in text


# ─── RTF ─────────────────────────────────────────────────────────────────────

def test_extract_rtf():
    from src.api.file_extractor import extract_text
    rtf_content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Arial;}} \f0\fs24 Texto em RTF.}"
    fs = FakeFileStorage(rtf_content.encode("latin-1"), "doc.rtf")
    text, err = extract_text(fs)
    assert err is None
    assert "RTF" in text or "Texto" in text


# ─── ODT ─────────────────────────────────────────────────────────────────────

def test_extract_odt():
    from src.api.file_extractor import extract_text
    from odf.opendocument import OpenDocumentText
    from odf.text import P
    from odf.element import Text

    doc = OpenDocumentText()
    p = P()
    p.addText("Texto extraído de arquivo ODT.")
    doc.text.addElement(p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fs = FakeFileStorage(buf.read(), "documento.odt")
    text, err = extract_text(fs)
    assert err is None
    assert "ODT" in text


# ─── Formato não suportado ────────────────────────────────────────────────────

def test_unsupported_format():
    from src.api.file_extractor import extract_text, SUPPORTED_EXTENSIONS
    # .xyz não está na lista oficial
    assert ".xyz" not in SUPPORTED_EXTENSIONS
    # Conteúdo binário puro com bytes não-imprimíveis → extrator retorna o que conseguir
    fs = FakeFileStorage(b"\x00\x01\x02binary", "arquivo.xyz")
    text, err = extract_text(fs)
    # Comportamento esperado: retorna conteúdo (como fallback) ou mensagem de erro
    assert isinstance(text, str) and isinstance((err or ""), str)


# ─── SUPPORTED_EXTENSIONS ────────────────────────────────────────────────────

def test_supported_extensions_contains_expected():
    from src.api.file_extractor import SUPPORTED_EXTENSIONS
    for ext in [".txt", ".pdf", ".docx", ".odt", ".rtf", ".csv", ".md", ".html"]:
        assert ext in SUPPORTED_EXTENSIONS, f"{ext} deveria estar em SUPPORTED_EXTENSIONS"
