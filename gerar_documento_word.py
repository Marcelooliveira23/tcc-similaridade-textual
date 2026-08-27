"""
Script para gerar documento Word (DOCX) conforme normas ABNT e template UNINTER.
Converte a dissertação de Markdown para formato Word profissional.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_margins(doc, top=3.0, bottom=2.0, left=3.0, right=2.0):
    """Define as margens do documento conforme ABNT."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top / 2.54)
        section.bottom_margin = Inches(bottom / 2.54)
        section.left_margin = Inches(left / 2.54)
        section.right_margin = Inches(right / 2.54)

def set_line_spacing(paragraph, spacing=1.5):
    """Define espaçamento entre linhas."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = spacing

def add_header_para(doc, text, level=1):
    """Adiciona parágrafo de seção com formatação adequada."""
    if level == 1:
        # Seção nível 1: MAIÚSCULAS COM NEGRITO
        p = doc.add_paragraph()
        p.style = 'Heading 1'
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        set_line_spacing(p, 1.5)
    elif level == 2:
        # Seção nível 2: MAIÚSCULAS SEM NEGRITO
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = False
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        set_line_spacing(p, 1.5)
    elif level == 3:
        # Seção nível 3: minúsculas com negrito
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        set_line_spacing(p, 1.5)
    
    # Espaçamento antes e depois
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    
    return p

def add_body_para(doc, text):
    """Adiciona parágrafo de corpo com formatação adequada."""
    p = doc.add_paragraph(text)
    p.style = 'Normal'
    
    # Formatação
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Arial'
    
    # Espaçamento
    set_line_spacing(p, 1.5)
    p.paragraph_format.first_line_indent = Inches(1.25 / 2.54)
    p.paragraph_format.space_after = Pt(0)
    
    return p

def create_tcc_document():
    """Cria o documento TCC conforme normas ABNT."""
    
    doc = Document()
    set_margins(doc, top=3.0, bottom=2.0, left=3.0, right=2.0)
    
    # PÁGINA DE TÍTULO
    title = doc.add_paragraph()
    title_run = title.add_run("SISTEMA DE COMPARAÇÃO DE SIMILARIDADE TEXTUAL:\nUMA ANÁLISE COMPARATIVA DE ALGORITMOS")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Arial'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_line_spacing(title, 1.5)
    title.paragraph_format.space_after = Pt(24)
    
    # Autor
    author_p = doc.add_paragraph()
    author_run = author_p.add_run("[Seu Nome Completo]")
    author_run.font.size = Pt(12)
    author_run.font.name = 'Arial'
    author_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_line_spacing(author_p, 1.5)
    author_p.paragraph_format.space_after = Pt(12)
    
    # Orientador
    prof_p = doc.add_paragraph()
    prof_run = prof_p.add_run("Prof. [Nome do Orientador]")
    prof_run.font.size = Pt(12)
    prof_run.font.name = 'Arial'
    prof_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_line_spacing(prof_p, 1.5)
    prof_p.paragraph_format.space_after = Pt(12)
    
    # Espaço
    doc.add_paragraph()
    
    # Instituição e data
    inst_p = doc.add_paragraph()
    inst_run = inst_p.add_run("Centro Universitário UNINTER\nEscola Superior Politécnica\n\n[Cidade], [Mês] de [Ano]")
    inst_run.font.size = Pt(12)
    inst_run.font.name = 'Arial'
    inst_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_line_spacing(inst_p, 1.5)
    
    # Quebra de página
    doc.add_page_break()
    
    # RESUMO
    add_header_para(doc, "RESUMO", level=1)
    
    resumo_text = """O presente trabalho apresenta uma análise comparativa de três algoritmos de similaridade textual: TF-IDF com Similaridade de Cosseno, Coeficiente de Jaccard e Distância de Levenshtein. A similaridade textual é um problema fundamental em diversas áreas da Computação, com aplicações em detecção de plágio, recuperação de informação e deduplicação de dados. Foi desenvolvido um sistema funcional em arquitetura em camadas que implementa os três algoritmos e oferece uma interface web para comparação de textos. A avaliação foi realizada sobre um dataset de 26 pares de textos rotulados, utilizando métricas de desempenho como F1-Score, Precision, Recall e Accuracy. Os resultados indicam que cada algoritmo possui forças e fraquezas distintas, sendo o TF-IDF mais apropriado para análise semântica, o Jaccard para comparações rápidas e o Levenshtein para detecção de erros ortográficos. Conclui-se que a escolha do algoritmo deve estar alinhada com o contexto de aplicação específico."""
    
    resumo_p = doc.add_paragraph(resumo_text)
    for run in resumo_p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Arial'
    set_line_spacing(resumo_p, 1.0)
    resumo_p.paragraph_format.space_after = Pt(0)
    
    # Palavras-chave
    kw_p = doc.add_paragraph()
    kw_p.add_run("Palavras-chave: ").bold = True
    kw_run = kw_p.add_run("similaridade textual, algoritmos de comparação, TF-IDF, Jaccard, Levenshtein, processamento de linguagem natural.")
    kw_run.font.size = Pt(10)
    kw_run.font.name = 'Arial'
    set_line_spacing(kw_p, 1.0)
    kw_p.paragraph_format.space_after = Pt(12)
    
    doc.add_page_break()
    
    # INTRODUÇÃO
    add_header_para(doc, "1 INTRODUÇÃO", level=1)
    
    add_header_para(doc, "1.1 Problemática", level=3)
    
    intro_text1 = """A detecção e quantificação de similaridade entre textos é um problema fundamental em diversas áreas da Ciência da Computação. Na era digital, com a crescente geração de dados textuais em redes sociais, plataformas de compartilhamento, bases de dados acadêmicas e sistemas corporativos, a automatização dessa tarefa tornou-se crítica para eficiência operacional e segurança da informação."""
    
    add_body_para(doc, intro_text1)
    
    intro_text2 = """Diferentes contextos de aplicação, porém, demandam diferentes abordagens. Enquanto alguns cenários exigem precisão elevada na comparação de palavras-chave—como na detecção de plágio acadêmico—outros requerem robustez a pequenas variações ortográficas e morfológicas, como no matching de nomes em sistemas de registros públicos. A literatura e a prática profissional apontam que não existe um algoritmo único que seja ótimo para todos os casos de uso."""
    
    add_body_para(doc, intro_text2)
    
    add_header_para(doc, "1.2 Questão de Pesquisa", level=3)
    
    quest_text = """Qual é o desempenho comparativo de três algoritmos clássicos de similaridade textual (TF-IDF com Similaridade de Cosseno, Coeficiente de Jaccard e Distância de Levenshtein) quando aplicados a cenários reais de comparação, e em quais contextos cada um se destaca?"""
    
    add_body_para(doc, quest_text)
    
    add_header_para(doc, "1.3 Justificativa", level=3)
    
    just_text = """Este trabalho é justificado por três razões principais. Primeiro, do ponto de vista acadêmico, a comparação sistemática de algoritmos clássicos de similaridade contribui para a formação de engenheiros de software capazes de tomar decisões informadas sobre qual ferramenta utilizar em cada contexto. Segundo, do ponto de vista prático, as conclusões deste trabalho informam decisões de arquitetura em sistemas reais que necessitam de busca, deduplicação, recomendação ou análise de similaridade. Terceiro, embora existam muitos estudos individuais de cada algoritmo, há uma lacuna relativa na literatura de trabalhos que os comparam de forma sistemática e rigorosa em um mesmo contexto de aplicação."""
    
    add_body_para(doc, just_text)
    
    add_header_para(doc, "1.4 Objetivos", level=3)
    
    add_header_para(doc, "1.4.1 Objetivo Geral", level=2)
    
    obj_geral_text = """Desenvolver um sistema funcional de comparação de similaridade textual que implemente três algoritmos distintos e realizar uma análise comparativa estruturada de seu desempenho em diferentes cenários de aplicação."""
    
    add_body_para(doc, obj_geral_text)
    
    add_header_para(doc, "1.4.2 Objetivos Específicos", level=2)
    
    # Lista de objetivos
    objectives = [
        "Implementar três algoritmos de similaridade textual em uma arquitetura em camadas (API, serviço, algoritmo, repositório);",
        "Disponibilizar uma interface web funcional para entrada de textos (colagem ou upload de arquivos) e comparação;",
        "Definir e executar uma metodologia de avaliação com métricas objetivas (F1-Score, Precision, Recall, Accuracy);",
        "Gerar um dataset de teste com 26 pares de textos manualmente rotulados como similares ou dissimilares;",
        "Comparar o desempenho dos três algoritmos e identificar os cenários onde cada um se destaca;",
        "Documentar conclusões e limitações de cada abordagem para orientar futuros desenvolvedores."
    ]
    
    for i, obj in enumerate(objectives, 1):
        obj_p = doc.add_paragraph(obj, style='List Number')
        for run in obj_p.runs:
            run.font.size = Pt(12)
            run.font.name = 'Arial'
        set_line_spacing(obj_p, 1.5)
        obj_p.paragraph_format.space_after = Pt(0)
    
    doc.add_page_break()
    
    # REFERENCIAL TEÓRICO
    add_header_para(doc, "2 REFERENCIAL TEÓRICO", level=1)
    
    # Seção 2.1
    add_header_para(doc, "2.1 Similaridade Textual e Aplicações", level=3)
    
    ref_text1 = """Similaridade textual é uma medida quantitativa de proximidade ou semelhança entre dois ou mais documentos de texto. Esta medida é fundamental para uma variedade de aplicações práticas, desde a recuperação de informação até a detecção de plágio. A formalização desta medida é essencial para a implementação de algoritmos que possam automatizar processos que dependem da comparação de textos."""
    
    add_body_para(doc, ref_text1)
    
    ref_text2 = """Na prática, a similaridade textual encontra aplicação em diversos domínios. Em sistemas de recuperação de informação, como motores de busca, a similaridade entre uma consulta e os documentos indexados é utilizada para rankear resultados. Em ambientes académicos e corporativos, a detecção de plágio utiliza similaridade textual para identificar possíveis cópias ou infrações de propriedade intelectual. Em e-commerce e plataformas de análise de dados, a deduplicação de registros e a análise de reviews dependem de métricas de similaridade. Finalmente, em sistemas de recomendação, textos similares são utilizados para sugerir conteúdo relevante aos usuários."""
    
    add_body_para(doc, ref_text2)
    
    doc.add_page_break()
    
    # Salvar documento
    output_path = r"C:\Users\mrced\OneDrive\Documents\TCC\TCC_FORMATADO.docx"
    doc.save(output_path)
    
    print(f"✅ Documento criado com sucesso: {output_path}")
    return output_path

if __name__ == "__main__":
    create_tcc_document()
    print("✅ Processo concluído!")
