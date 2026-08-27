import docx
import os

doc_path = "TCC_FINAL.docx"
if not os.path.exists(doc_path):
    print(f"Erro: {doc_path} nao existe.")
    exit(1)

doc = docx.Document(doc_path)
text_lines = []
for p in doc.paragraphs:
    val = p.text.strip()
    if val:
        text_lines.append(val)

# Also look into tables text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            val = cell.text.strip()
            if val:
                # Add individual cell elements if not already there, but or just join standard paragraphs
                pass

full_text = "\n".join(text_lines)

terms = ['TF-IDF', 'Jaccard', 'Levenshtein', 'Accuracy', 'Precision', 'Recall', 'F1', '/compare', '/evaluate', 'dataset', 'base_pairs']
print("==== RESUMO BOOLEANO POR TERMO ====")
for t in terms:
    presence = t in full_text
    print(f"{t}: {presence}")

print("\n==== 20 PRIMEIRAS LINHAS UTEIS ====")
for i, line in enumerate(text_lines[:20]):
    print(f"{i+1:02d}: {line}")