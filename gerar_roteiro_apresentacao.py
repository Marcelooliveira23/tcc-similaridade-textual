"""
Script para gerar documento de ROTEIRO DE APRESENTAÇÃO ORAL em Word (DOCX).
Converte a defesa de Markdown para formato Word profissional para apresentação.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_margins(doc, top=3.0, bottom=2.0, left=3.0, right=2.0):
    """Define as margens conforme ABNT."""
    for section in doc.sections:
        section.top_margin = Inches(top / 2.54)
        section.bottom_margin = Inches(bottom / 2.54)
        section.left_margin = Inches(left / 2.54)
        section.right_margin = Inches(right / 2.54)

def set_line_spacing(paragraph, spacing=1.5):
    """Define espaçamento entre linhas."""
    paragraph.paragraph_format.line_spacing = spacing

def create_presentation():
    """Cria o documento de apresentação para defesa."""
    
    doc = Document()
    set_margins(doc)
    
    # CAPA
    title = doc.add_paragraph()
    title_run = title.add_run("ROTEIRO DE APRESENTAÇÃO PARA DEFESA")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Arial'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_line_spacing(title, 1.5)
    title.paragraph_format.space_after = Pt(24)
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Sistema de Comparação de Similaridade Textual: Uma Análise Comparativa de Algoritmos")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.name = 'Arial'
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_line_spacing(subtitle, 1.5)
    subtitle.paragraph_format.space_after = Pt(24)
    
    time_p = doc.add_paragraph()
    time_run = time_p.add_run("Duração Total: 15 minutos\nData: [Data da Defesa]\nLocal: [Local da Defesa]")
    time_run.font.size = Pt(12)
    time_run.font.name = 'Arial'
    time_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_line_spacing(time_p, 1.5)
    time_p.paragraph_format.space_after = Pt(24)
    
    doc.add_page_break()
    
    # SLIDE 1
    slide1 = doc.add_paragraph()
    slide1_run = slide1.add_run("SLIDE 1: TÍTULO E APRESENTAÇÃO PESSOAL")
    slide1_run.bold = True
    slide1_run.font.size = Pt(12)
    slide1_run.font.name = 'Arial'
    set_line_spacing(slide1, 1.5)
    
    timing1 = doc.add_paragraph()
    timing1_run = timing1.add_run("⏱️ Duração: 1 minuto")
    timing1_run.bold = True
    timing1_run.font.size = Pt(11)
    timing1_run.font.name = 'Arial'
    timing1_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing1, 1.5)
    
    content1 = doc.add_paragraph("\"Bom [manhã/tarde]. Meu nome é [Seu Nome], sou aluno(a) de Engenharia de Software da UNINTER. Hoje vou apresentar meu Trabalho de Conclusão de Curso com o tema 'Sistema de Comparação de Similaridade Textual: Uma Análise Comparativa de Algoritmos'. Este trabalho foi desenvolvido sob orientação do Prof. [Nome do Orientador].\"")
    for run in content1.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(content1, 1.5)
    
    doc.add_paragraph()  # Espaço
    
    # SLIDE 2
    slide2 = doc.add_paragraph()
    slide2_run = slide2.add_run("SLIDE 2: O PROBLEMA")
    slide2_run.bold = True
    slide2_run.font.size = Pt(12)
    slide2_run.font.name = 'Arial'
    set_line_spacing(slide2, 1.5)
    
    timing2 = doc.add_paragraph()
    timing2_run = timing2.add_run("⏱️ Duração: 1.5 minutos")
    timing2_run.bold = True
    timing2_run.font.size = Pt(11)
    timing2_run.font.name = 'Arial'
    timing2_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing2, 1.5)
    
    content2 = doc.add_paragraph("\"Na era digital, bilhões de documentos são criados diariamente. Empresas, universidades e governos enfrentam o desafio de comparar textos: detectar plágio, eliminar duplicatas, recomendar conteúdo similar. Cada algoritmo de similaridade tem forças e fraquezas diferentes. A pergunta é: qual algoritmo escolher para cada situação?\"")
    for run in content2.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(content2, 1.5)
    
    context = doc.add_paragraph("Dados de Contexto: Mais de 5 milhões de artigos científicos publicados anualmente; empresas como Google, Facebook e Microsoft usam similaridade textual em seus serviços; detecção de plágio é exigência em instituições educacionais.")
    for run in context.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True
    set_line_spacing(context, 1.5)
    
    doc.add_paragraph()  # Espaço
    
    # SLIDE 3
    slide3 = doc.add_paragraph()
    slide3_run = slide3.add_run("SLIDE 3: OBJETIVOS")
    slide3_run.bold = True
    slide3_run.font.size = Pt(12)
    slide3_run.font.name = 'Arial'
    set_line_spacing(slide3, 1.5)
    
    timing3 = doc.add_paragraph()
    timing3_run = timing3.add_run("⏱️ Duração: 1 minuto")
    timing3_run.bold = True
    timing3_run.font.size = Pt(11)
    timing3_run.font.name = 'Arial'
    timing3_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing3, 1.5)
    
    obj_intro = doc.add_paragraph("\"Os objetivos deste trabalho são:\"")
    for run in obj_intro.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(obj_intro, 1.5)
    
    objectives = [
        "Implementar três algoritmos clássicos: TF-IDF, Jaccard e Levenshtein;",
        "Criar uma interface web funcional para testar os algoritmos;",
        "Estabelecer metodologia rigorosa de avaliação com métricas objetivas;",
        "Comparar o desempenho dos três algoritmos em diferentes cenários;",
        "Definir recomendações de uso para cada algoritmo."
    ]
    
    for obj in objectives:
        obj_p = doc.add_paragraph(obj, style='List Bullet')
        for run in obj_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(obj_p, 1.5)
    
    doc.add_page_break()
    
    # SLIDE 4
    slide4 = doc.add_paragraph()
    slide4_run = slide4.add_run("SLIDE 4: ARQUITETURA DO SISTEMA")
    slide4_run.bold = True
    slide4_run.font.size = Pt(12)
    slide4_run.font.name = 'Arial'
    set_line_spacing(slide4, 1.5)
    
    timing4 = doc.add_paragraph()
    timing4_run = timing4.add_run("⏱️ Duração: 2 minutos")
    timing4_run.bold = True
    timing4_run.font.size = Pt(11)
    timing4_run.font.name = 'Arial'
    timing4_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing4, 1.5)
    
    arch_intro = doc.add_paragraph("\"O sistema segue uma arquitetura profissional em camadas, o que permite modularidade e facilita a manutenção:\"")
    for run in arch_intro.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(arch_intro, 1.5)
    
    layers = [
        "Camada de Apresentação: Interface web em HTML5, CSS3 e JavaScript para o usuário inserir textos ou fazer upload de arquivos;",
        "Camada de API: REST endpoints implementados em Flask Python, fornecendo endpoints para /compare, /evaluate, /report;",
        "Camada de Serviço: ComparisonService que orquestra os algoritmos e gerencia fluxo de negócio;",
        "Camada de Algoritmos: Implementação isolada dos três algoritmos com testes unitários;",
        "Camada de Dados: Persistência em SQLite com padrão Repository para abstração."
    ]
    
    for layer in layers:
        layer_p = doc.add_paragraph(layer, style='List Bullet')
        for run in layer_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(layer_p, 1.5)
    
    padroes = doc.add_paragraph("Padrões de Design: Repository Pattern, Strategy Pattern, Service Layer Pattern, Decorator Pattern para pipeline de pré-processamento.")
    for run in padroes.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True
    set_line_spacing(padroes, 1.5)
    
    doc.add_paragraph()  # Espaço
    
    # SLIDE 5
    slide5 = doc.add_paragraph()
    slide5_run = slide5.add_run("SLIDE 5: ALGORITMO 1 - TF-IDF COM SIMILARIDADE DE COSSENO")
    slide5_run.bold = True
    slide5_run.font.size = Pt(12)
    slide5_run.font.name = 'Arial'
    set_line_spacing(slide5, 1.5)
    
    timing5 = doc.add_paragraph()
    timing5_run = timing5.add_run("⏱️ Duração: 1.5 minutos")
    timing5_run.bold = True
    timing5_run.font.size = Pt(11)
    timing5_run.font.name = 'Arial'
    timing5_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing5, 1.5)
    
    tfidf_desc = doc.add_paragraph("\"TF-IDF significa Term Frequency-Inverse Document Frequency. É uma abordagem vetorial que representa cada documento como um vetor numérico onde cada dimensão é uma palavra e o valor é o peso de importância daquela palavra para o documento.\"")
    for run in tfidf_desc.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(tfidf_desc, 1.5)
    
    tfidf_strengths = [
        "✅ Robusto em textos longos e semanticamente ricos;",
        "✅ Detecta similaridade conceitual, não apenas lexical;",
        "✅ Funciona bem em contextos de busca e recomendação;",
        "✅ Amplitude de aplicação bem estabelecida na indústria."
    ]
    
    for strength in tfidf_strengths:
        s_p = doc.add_paragraph(strength, style='List Bullet')
        for run in s_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(s_p, 1.5)
    
    tfidf_weaknesses = [
        "❌ Ignora completamente a ordem das palavras;",
        "❌ Computacionalmente mais caro que alternativas;",
        "❌ Pior desempenho em textos muito curtos;",
        "❌ Menos eficaz em detectar erros ortográficos."
    ]
    
    for weak in tfidf_weaknesses:
        w_p = doc.add_paragraph(weak, style='List Bullet')
        for run in w_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(w_p, 1.5)
    
    doc.add_paragraph()  # Espaço
    
    # SLIDE 6
    slide6 = doc.add_paragraph()
    slide6_run = slide6.add_run("SLIDE 6: ALGORITMO 2 - COEFICIENTE DE JACCARD")
    slide6_run.bold = True
    slide6_run.font.size = Pt(12)
    slide6_run.font.name = 'Arial'
    set_line_spacing(slide6, 1.5)
    
    timing6 = doc.add_paragraph()
    timing6_run = timing6.add_run("⏱️ Duração: 1.5 minutos")
    timing6_run.bold = True
    timing6_run.font.size = Pt(11)
    timing6_run.font.name = 'Arial'
    timing6_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing6, 1.5)
    
    jaccard_desc = doc.add_paragraph("\"O Jaccard é uma medida baseada em teoria dos conjuntos. Calcula a razão entre o tamanho da interseção de dois conjuntos e o tamanho da sua união. Para textos, os conjuntos são as palavras únicas de cada documento.\"")
    for run in jaccard_desc.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(jaccard_desc, 1.5)
    
    jaccard_strengths = [
        "✅ Muito rápido computacionalmente;",
        "✅ Determinístico—sempre produz o mesmo resultado;",
        "✅ Ideal para deduplicação em larga escala;",
        "✅ Simples de implementar e compreender."
    ]
    
    for strength in jaccard_strengths:
        s_p = doc.add_paragraph(strength, style='List Bullet')
        for run in s_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(s_p, 1.5)
    
    jaccard_weaknesses = [
        "❌ Não captura semântica ou conceitos;",
        "❌ Falha com textos muito curtos;",
        "❌ Sensível a variações menores do texto;",
        "❌ Não considera frequência de palavras."
    ]
    
    for weak in jaccard_weaknesses:
        w_p = doc.add_paragraph(weak, style='List Bullet')
        for run in w_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(w_p, 1.5)
    
    doc.add_page_break()
    
    # SLIDE 7
    slide7 = doc.add_paragraph()
    slide7_run = slide7.add_run("SLIDE 7: ALGORITMO 3 - DISTÂNCIA DE LEVENSHTEIN")
    slide7_run.bold = True
    slide7_run.font.size = Pt(12)
    slide7_run.font.name = 'Arial'
    set_line_spacing(slide7, 1.5)
    
    timing7 = doc.add_paragraph()
    timing7_run = timing7.add_run("⏱️ Duração: 1.5 minutos")
    timing7_run.bold = True
    timing7_run.font.size = Pt(11)
    timing7_run.font.name = 'Arial'
    timing7_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing7, 1.5)
    
    lev_desc = doc.add_paragraph("\"Levenshtein calcula o número mínimo de edições de caracteres (inserção, deleção, substituição) necessárias para transformar uma string em outra. É baseado em programação dinâmica e é especialmente útil para detecção de erros de digitação.\"")
    for run in lev_desc.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(lev_desc, 1.5)
    
    lev_strengths = [
        "✅ Excelente para detecção de typos e erros ortográficos;",
        "✅ Intuitivo de compreender—mostra diferenças de caracteres;",
        "✅ Funciona bem para strings muito curtas;",
        "✅ Aplicável em contextos de correção automática."
    ]
    
    for strength in lev_strengths:
        s_p = doc.add_paragraph(strength, style='List Bullet')
        for run in s_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(s_p, 1.5)
    
    lev_weaknesses = [
        "❌ Computacionalmente caro para strings longas;",
        "❌ Não captura semântica;",
        "❌ Não considera contexto ou significado;",
        "❌ Menos eficaz em paráfrases e reformulações."
    ]
    
    for weak in lev_weaknesses:
        w_p = doc.add_paragraph(weak, style='List Bullet')
        for run in w_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(w_p, 1.5)
    
    doc.add_paragraph()  # Espaço
    
    # SLIDE 8
    slide8 = doc.add_paragraph()
    slide8_run = slide8.add_run("SLIDE 8: METODOLOGIA DE AVALIAÇÃO")
    slide8_run.bold = True
    slide8_run.font.size = Pt(12)
    slide8_run.font.name = 'Arial'
    set_line_spacing(slide8, 1.5)
    
    timing8 = doc.add_paragraph()
    timing8_run = timing8.add_run("⏱️ Duração: 1.5 minutos")
    timing8_run.bold = True
    timing8_run.font.size = Pt(11)
    timing8_run.font.name = 'Arial'
    timing8_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing8, 1.5)
    
    method_intro = doc.add_paragraph("\"Para comparar os algoritmos de forma justa, utilizamos uma metodologia rigorosa:\"")
    for run in method_intro.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(method_intro, 1.5)
    
    methodology = [
        "Dataset: 26 pares de textos em português, manualmente rotulados como similares (20) ou dissimilares (6);",
        "Métricas: F1-Score, Precision, Recall, Accuracy—todas baseadas em matriz de confusão;",
        "Threshold: 0.5—scores acima de 0.5 são classificados como similares, abaixo como dissimilares;",
        "Cenários: Textos paráfraseados, com erros ortográficos, duplicados, completamente não relacionados."
    ]
    
    for method in methodology:
        m_p = doc.add_paragraph(method, style='List Bullet')
        for run in m_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(m_p, 1.5)
    
    dataset_note = doc.add_paragraph("Nota: O dataset foi cuidadosamente curado para representar cenários reais que encontramos em aplicações práticas.")
    for run in dataset_note.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True
    set_line_spacing(dataset_note, 1.5)
    
    doc.add_page_break()
    
    # SLIDE 9
    slide9 = doc.add_paragraph()
    slide9_run = slide9.add_run("SLIDE 9: RESULTADOS COMPARATIVOS")
    slide9_run.bold = True
    slide9_run.font.size = Pt(12)
    slide9_run.font.name = 'Arial'
    set_line_spacing(slide9, 1.5)
    
    timing9 = doc.add_paragraph()
    timing9_run = timing9.add_run("⏱️ Duração: 1.5 minutos")
    timing9_run.bold = True
    timing9_run.font.size = Pt(11)
    timing9_run.font.name = 'Arial'
    timing9_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing9, 1.5)
    
    results_intro = doc.add_paragraph("\"Os resultados mostram que cada algoritmo tem um perfil diferente de desempenho:\"")
    for run in results_intro.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    set_line_spacing(results_intro, 1.5)
    
    results_list = [
        "TF-IDF: Accuracy 84.6% - Melhor em análise semântica, detecta textos paráfraseados;",
        "Jaccard: Accuracy 80.8% - Melhor em deduplicação rápida, determinístico;",
        "Levenshtein: Accuracy 73.1% - Melhor em erros ortográficos, variações léxicas."
    ]
    
    for result in results_list:
        r_p = doc.add_paragraph(result, style='List Bullet')
        for run in r_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(r_p, 1.5)
    
    conclusion_note = doc.add_paragraph("Conclusão: Não há um algoritmo universalmente melhor. A escolha deve considerar o contexto de aplicação específico.")
    for run in conclusion_note.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.bold = True
    set_line_spacing(conclusion_note, 1.5)
    
    doc.add_paragraph()  # Espaço
    
    # SLIDE 10
    slide10 = doc.add_paragraph()
    slide10_run = slide10.add_run("SLIDE 10: RECOMENDAÇÕES E CONCLUSÕES")
    slide10_run.bold = True
    slide10_run.font.size = Pt(12)
    slide10_run.font.name = 'Arial'
    set_line_spacing(slide10, 1.5)
    
    timing10 = doc.add_paragraph()
    timing10_run = timing10.add_run("⏱️ Duração: 2 minutos")
    timing10_run.bold = True
    timing10_run.font.size = Pt(11)
    timing10_run.font.name = 'Arial'
    timing10_run.font.color.rgb = RGBColor(0, 0, 255)
    set_line_spacing(timing10, 1.5)
    
    recommendations = [
        "Para Detecção de Plágio: Use TF-IDF—é o mais eficaz em detectar similaridade semântica;",
        "Para Deduplicação: Use Jaccard—é rápido e determinístico, ideal em larga escala;",
        "Para Correção Ortográfica: Use Levenshtein—é especializado em erros de caracteres;",
        "Para Sistemas Híbridos: Combine os três—cada um fornece uma perspectiva diferente."
    ]
    
    for rec in recommendations:
        rec_p = doc.add_paragraph(rec, style='List Bullet')
        for run in rec_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(rec_p, 1.5)
    
    final_conclusion = doc.add_paragraph()
    final_run = final_conclusion.add_run("\"Em conclusão, este trabalho demonstrou que a seleção apropriada de algoritmo de similaridade textual é crítica para o sucesso de aplicações reais. A análise comparativa rigorosa fornece diretrizes práticas para engenheiros de software.\"")
    final_run.font.size = Pt(11)
    final_run.font.name = 'Arial'
    final_run.italic = True
    set_line_spacing(final_conclusion, 1.5)
    
    doc.add_page_break()
    
    # PERGUNTAS E RESPOSTAS
    qa_title = doc.add_paragraph()
    qa_title_run = qa_title.add_run("POSSÍVEIS PERGUNTAS E RESPOSTAS")
    qa_title_run.bold = True
    qa_title_run.font.size = Pt(14)
    qa_title_run.font.name = 'Arial'
    set_line_spacing(qa_title, 1.5)
    qa_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    qa_items = [
        ("P: Por que você escolheu estes três algoritmos?", "R: São algoritmos clássicos, bem estabelecidos, representam três paradigmas diferentes (vetorial, conjuntista, caractere-level) e são amplamente usados na indústria."),
        ("P: Por que não usou algoritmos mais sofisticados como Word2Vec ou BERT?", "R: O escopo era analisar algoritmos clássicos que são fundamentais. Algoritmos baseados em deep learning são mais complexos e seria necessário um trabalho separado."),
        ("P: Como você se certificou de que o dataset é representativo?", "R: Incluímos manualmente pares de diferentes tipos: paráfrases, erros ortográficos, duplicatas, textos não relacionados. Isso reflete cenários reais."),
        ("P: Por que threshold 0.5?", "R: Foi determinado empiricamente para balancear os algoritmos. Diferentes aplicações podem usar diferentes thresholds dependendo de suas prioridades."),
        ("P: O sistema está pronto para produção?", "R: O código está bem estruturado, testado e pode ser usado como base. Seria necessário otimizações de performance para volumes muito grandes de dados."),
        ("P: Como o sistema lida com outros idiomas?", "R: A prototipação foi feita em português. A generalização para outros idiomas requer adaptação das regras de tokenização e stopwords."),
    ]
    
    for i, (pergunta, resposta) in enumerate(qa_items, 1):
        doc.add_paragraph()
        p_p = doc.add_paragraph(pergunta)
        for run in p_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.bold = True
        set_line_spacing(p_p, 1.5)
        
        r_p = doc.add_paragraph(resposta)
        for run in r_p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        set_line_spacing(r_p, 1.5)
    
    # Salvar documento
    output_path = r"C:\Users\mrced\OneDrive\Documents\TCC\ROTEIRO_APRESENTACAO_DEFESA.docx"
    doc.save(output_path)
    
    print(f"✅ Roteiro de apresentação criado: {output_path}")
    print(f"✅ 10 slides com timings e conteúdo completo")
    print(f"✅ Possíveis perguntas e respostas inclusas")
    print(f"✅ Pronto para apresentação oral de 15 minutos")
    
    return output_path

if __name__ == "__main__":
    create_presentation()
    print("\n✅ Processo concluído com sucesso!")
